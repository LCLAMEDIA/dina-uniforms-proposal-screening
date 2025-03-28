import os
from docx import Document
import json
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed

from VoiceflowOperations import VoiceflowOperations
from GoogleDocsOperations import GoogleDocsOperations
from PromptsOperations import PromptsOperations
from GPTOperations import GPTOperations
from NotionOperator import NotionOperator
from DocxOperator import DocxOperator
from Analysis import Analysis
import pandas as pd
import io
from typing import Optional

logging.basicConfig(level=logging.INFO)

class DocumentContentExtractor:
    def __init__(self, document_path: Optional[str] = None, document_bytes: Optional[bytes] = None):
        """
        Accepts either document from path or document as bytes.
        """

        document_data = None
        
        if isinstance(document_path, str):
            document_data = document_path

        if isinstance(document_bytes, bytes):
            document_data = io.BytesIO(document_bytes)
        
        if not document_data:
            raise Exception('Document data not populated')
        
        self.doc = Document(document_data)
        logging.info(self.doc.element.body)
        
    def extract_content(self):
        content_parts = []
        dataframes = []  # List to store DataFrames
        logging.info("Starting content extraction.")
        for element in self.doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                para = [e for e in self.doc.paragraphs if e._element is element][0]
                text = para.text
                if text:  # Ensure the paragraph contains text
                    content_parts.append(text)
            elif element.tag.endswith('tbl'):  # Table
                table = [t for t in self.doc.tables if t._element is element][0]
                if table:
                    table_data = self._table_to_json(table)
                    if table_data:
                        df = pd.DataFrame(table_data)
                        dataframes.append(df)
                        table_str = self._table_data_to_string(table_data)
                        content_parts.append(table_str)
                    else:
                        logging.info("Empty table encountered, skipping")
        
         # Print each DataFrame in the array
        for i, df in enumerate(dataframes):
            print(f"DataFrame {i+1}:\n{df}\n")
                 
        # Save each DataFrame to a separate sheet in an Excel file
        # with pd.ExcelWriter('DataFrames.xlsx') as writer:
        #     for i, df in enumerate(dataframes):
        #         df.to_excel(writer, sheet_name=f'DataFrame {i+1}')
        #         print(f"DataFrame {i+1}:\n{df}\n")
        
        # Join all parts into one flattened string
        return '\n'.join(content_parts)


    def _table_to_json(self, table):
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        table_json = []
        for row in table.rows[1:]:
            logging.error(headers)
            logging.error(row.cells)
            row_data = {headers[i]: cell.text.strip() for i, cell in enumerate(row.cells) if i + 1 <= len(headers)}
            table_json.append(row_data)
        return table_json

    def _table_data_to_string(self, table_data):
        # Convert each row dictionary to a string and join all with newline
        return '\n'.join([json.dumps(row) for row in table_data])
    
    
class ProposalScreeningOperations:
    def __init__(
        self,
        proposal_url: str,
        google_docs_ops: GoogleDocsOperations,
        voiceflow_ops: VoiceflowOperations,
        prompts_ops: PromptsOperations,
        gpt_ops: GPTOperations,
        notion_ops: NotionOperator,
        docx_ops: DocxOperator,
        page_id: str
    ):
        self.proposal_url = proposal_url
        self.google_docs_ops = google_docs_ops
        self.voiceflow_ops = voiceflow_ops
        self.prompt_ops = prompts_ops
        self.gpt_ops = gpt_ops
        self.notion_ops = notion_ops
        self.docx_ops = docx_ops
        self.page_id = page_id

    def split_into_chunks(
        self, text, chunk_size: int = 8000, overlap_percentage: float = 0.1
    ) -> list[str]:
        # Calculate the overlap in terms of characters
        overlap = int(chunk_size * overlap_percentage)

        # Initialize the list to store chunks
        chunks = []

        # Calculate the start and end indices for each chunk and extract the chunks
        start_index = 0
        while start_index < len(text):
            # Ensure the last chunk includes the end of the text
            end_index = min(start_index + chunk_size, len(text))
            chunk = text[start_index:end_index]
            chunks.append(chunk)

            # Move to the next chunk, ensuring to overlap as specified
            start_index += chunk_size - overlap

        return chunks

    def extract_text(self, document_path: Optional[str] = None, document_bytes: Optional[bytes] = None, document_filename: str = ''):
        # Extract text from downloaded document
        if document_filename.endswith("docx"):
            logging.info("[Extract Text] Using DocumentExtractor")
            content = DocumentContentExtractor(document_path=document_path, document_bytes=document_bytes).extract_content()
        elif document_filename.endswith("pdf"):
            pass
        else:
            pass
        
        logging.info('[ProposalScreeningOperations] Extracted Text')
        return content

    def download_file(self, document_url: str, output_path: str):
        import requests
        
        # Send a GET request to the URL
        response = requests.get(document_url)
        
        # Check if the request was successful
        if response.status_code == 200:
            # Open the file in binary write mode and write the content
            with open(output_path, 'wb') as file:
                file.write(response.content)
        else:
            # Handle possible errors
            response.raise_for_status()
        
        return output_path


    def analyse_single_prompt(self, chunk: str, prompt_function) -> Analysis:
        """Run a single prompt for a single chunk, used concurrently within the chunk processor

        Args:
            chunk (str): Chunk of text we are running on
            prompt_function (function): Function that gets the prompt we want

        Returns:
            Analysis: _description_
        """
        prompt = prompt_function()
        raw = self.gpt_ops.query_chatgpt(
            f"{prompt.get('prompt')} Proposal Extract: {chunk}"
        )
        parsed = self.gpt_ops.parse_json_response(raw)
        return Analysis(chunk, prompt, parsed)
    
    def analyse_single_chunk(self, chunk: str) -> list[Analysis]:
        """Run All prompts and questions/checks for this single chunk

        Args:
            chunk (str): _description_
        """
        analysis_list = []
        with ThreadPoolExecutor(max_workers=15) as executor:
            # Submit all tasks to the executor
            futures = [executor.submit(self.analyse_single_prompt, chunk, prompt_function) for prompt_function in self.prompt_ops.all_prompts]

            # Optionally, wait for all futures to complete and handle any exceptions
            for future in as_completed(futures):
                try:
                    # Result method would raise any exceptions caught during the execution of the task
                    analysis_result = future.result()
                    analysis_list.append(analysis_result)
                except Exception as e:
                    logging.error(f"Error processing single prompt in chunk: {e}")

        return analysis_list

    def analyse_all_chunks(self, chunks: list[str]) -> list[Analysis]:
        """Loops over each text chunk, calls analyse_single_chunk, appends output to response

        Args:
            chunks (list[str]): List of chunked up proposal
        """
        analysis_list = []
        with ThreadPoolExecutor(max_workers=15) as executor:
            # Submit all tasks to the executor
            futures = [executor.submit(self.analyse_single_chunk, chunk) for chunk in chunks]

            # Optionally, wait for all futures to complete and handle any exceptions
            for future in as_completed(futures):
                try:
                    # Result method would raise any exceptions caught during the execution of the task
                    analysis_result = future.result()
                    analysis_list.append(analysis_result)
                except Exception as e:
                    logging.error(f"Error processing chunk: {e}")

        return analysis_list
    
    def handle_table_prompts(self, key, value):
        tables = [table.response.get('table') for table in value]
        analysis_texts = [table.response.get('analysis') for table in value if 'analysis' in table.response]

        print(f"handle_table_prompts: {tables}")
        print(f"analysis_texts: {analysis_texts}")
        
        # Fetch Prompts
        combine_table_prompt = self.prompt_ops.combine_table_prompt()

        # Generate Combined Tables
        combined_tables_response = self.gpt_ops.query_chatgpt(
            f"{combine_table_prompt.get('prompt')} Tables: {json.dumps(tables)}"
        )
        combined_tables = self.gpt_ops.parse_json_response(combined_tables_response)

        # Generate Combined Analysis
        combined_analysis = {}
        if analysis_texts:
            combine_analysis_prompt = self.prompt_ops.combine_analysis_prompt()
            combined_analysis_response = self.gpt_ops.query_chatgpt(
                f"{combine_analysis_prompt.get('prompt')} Analysis: {' '.join(analysis_texts)}"
            )
            combined_analysis = self.gpt_ops.parse_json_response(combined_analysis_response)
        
        combined_output = {**combined_tables, **combined_analysis}

        # Fetch the prompt object from the mapping for output
        prompt_obj = self.prompt_ops.prompt_mapping.get(key)()
        return Analysis('', prompt_obj, combined_output)

    def handle_dot_point_analysis_prompts(self, key, value):
        analysis_text = '\n[Extract]'.join([analysis.response.get('analysis') for analysis in value])
        analysis_dot_point_summary = json.dumps([analysis.response.get('dot_point_summary') for analysis in value])
        
        # Fetch Prompts
        analysis_prompt = self.prompt_ops.combine_analysis_prompt()
        dot_point_prompt = self.prompt_ops.combine_dot_point_prompt()
        
        # Generate Combined Analysis
        analysis_combined = self.gpt_ops.parse_json_response(self.gpt_ops.query_chatgpt(
            f"{analysis_prompt.get('prompt')} Analysis: {analysis_text}"
        ))
        dot_point_combined = self.gpt_ops.parse_json_response(self.gpt_ops.query_chatgpt(
            f"{dot_point_prompt.get('prompt')} Dot Point Analysis: {analysis_dot_point_summary}"
        ))
        
        # Fetch the prompt object from the mapping for output
        prompt_obj = self.prompt_ops.prompt_mapping.get(key)() 
        return Analysis('', prompt_obj, analysis_combined | dot_point_combined)
    
    def handle_timelines_prompts(self, key, value):
        timelines = json.dumps([timeline.response.get('timeline') for timeline in value])
        # Fetch Prompts
        combine_timelines_prompt = self.prompt_ops.combine_timelines_prompt()
        
        # Generate Combined Analysis
        timelines_combined = self.gpt_ops.parse_json_response(self.gpt_ops.query_chatgpt(
            f"{combine_timelines_prompt.get('prompt')} Timeline: {timelines}"
        ))
        
        # Fetch the prompt object from the mapping for output
        prompt_obj = self.prompt_ops.prompt_mapping.get(key)() 
        return Analysis('', prompt_obj, timelines_combined)

    def handle_cost_value_prompts(self, key, value):
        cost_value = json.dumps([cost_value.response.get('cost_value') for cost_value in value])
        # Fetch Prompts
        combine_cost_values_prompt = self.prompt_ops.combine_cost_value_prompt()
        
        # Generate Combined Analysis
        cost_values_combined = self.gpt_ops.parse_json_response(self.gpt_ops.query_chatgpt(
            f"{combine_cost_values_prompt.get('prompt')} cost_value: {cost_value}"
        ))
        
        # Fetch the prompt object from the mapping for output
        prompt_obj = self.prompt_ops.prompt_mapping.get(key)() 
        return Analysis('', prompt_obj, cost_values_combined)
    
    def handle_combining_chunk_analysis(self, key, value):
        dot_point_analysis_prompts = ['in_person_requirements_prompt', 'eligibility_prompt', 'uniform_specification_prompt', 'customer_support_service_prompt', 'long_term_partnership_potential_prompt', 'risk_management_analysis_prompt', 'compliance_evaluation_prompt']
        timeline_prompts = ['timelines_prompt']
        cost_value_prompts = ['cost_value_prompt']
        table_prompts = ['table_prompt']
        if key in dot_point_analysis_prompts:
            analysis_obj = self.handle_dot_point_analysis_prompts(key,value)
        elif key in timeline_prompts:
            analysis_obj = self.handle_timelines_prompts(key, value)
        elif key in cost_value_prompts:
            analysis_obj = self.handle_cost_value_prompts(key, value)
        elif key in table_prompts:
            analysis_obj = self.handle_table_prompts(key, value)

        
        return analysis_obj
    
    def combine_chunked_analysis(self, analysis_list: list[Analysis]):
        # Loop over all chunks, concatenating their analysis by prompt
        analysis_by_prompt = {}
        for chunk_analysis in analysis_list:
            for single_prompt_analysis in chunk_analysis:
                if single_prompt_analysis.prompt_name not in analysis_by_prompt:
                    analysis_by_prompt[single_prompt_analysis.prompt_name] = [single_prompt_analysis]
                else:
                    analysis_by_prompt[single_prompt_analysis.prompt_name] = analysis_by_prompt[single_prompt_analysis.prompt_name] + [single_prompt_analysis]
                    
        all_analysis = []
        with ThreadPoolExecutor(max_workers=15) as executor:
            # Submit all tasks to the executor
            futures = [executor.submit(self.handle_combining_chunk_analysis,key, value) for key, value in analysis_by_prompt.items()]

            # Optionally, wait for all futures to complete and handle any exceptions
            for future in as_completed(futures):
                try:
                    # Result method would raise any exceptions caught during the execution of the task
                    analysis_result = future.result()
                    all_analysis.append(analysis_result)
                except Exception as e:
                    logging.error(f"Error processing single prompt in chunk: {e}")
        
        return all_analysis

    
    def run(self):
        proposal_name = "Proposal"
        logging.info("Downloading File")
        file_location = self.download_file(self.proposal_url,'proposal.docx')
        logging.info("File Downloaded")
        
        # Extract text from proposal
        text = self.extract_text(file_location)

        # Split into chunks to feed into AI
        chunks = self.split_into_chunks(text, chunk_size=16000)
    
        # Loop over all chunks and generate an analysis for each
        analysis_list = self.analyse_all_chunks(chunks)
            
        combined_analysis_list = self.combine_chunked_analysis(analysis_list)
        
        for analysis in combined_analysis_list:
            print(analysis)
        
        
        self.notion_ops.create_page_from_analysis(proposal_name=proposal_name, analysis_list=combined_analysis_list, page_id=self.page_id)
        # Create Report
            
        # Analysis_list = [
        #     {
        #         "chunk_num": 1,
        #         "text": "Chunked Text"
        #         "analysis": [
        #             "prompt": "eligibility_prompt",
        #             "analysis": {
        #                 "GPT Fields"
        #             }
        #         ]
        #     }
        # ]

    def run_analysis_from_sharepoint(
            self,
            document_bytes: bytes,
            document_filename: str
        ):
        filename_without_extension = os.path.splitext(document_filename)[0]
        
        # Extract text from proposal
        text = self.extract_text(document_bytes=document_bytes, document_filename=document_filename)

        # Split into chunks to feed into AI
        chunks = self.split_into_chunks(text, chunk_size=16000)
    
        # Loop over all chunks and generate an analysis for each
        analysis_list = self.analyse_all_chunks(chunks)
            
        combined_analysis_list = self.combine_chunked_analysis(analysis_list)
        
        for analysis in combined_analysis_list:
            print(analysis)

        return self.docx_ops.create_page_from_analysis(proposal_name=filename_without_extension, analysis_list=combined_analysis_list, page_id=self.page_id)
        
        
        