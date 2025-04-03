from docx import Document
from datetime import datetime
import logging
from Analysis import Analysis
import pytz
import io

class DocxOperator:
    def __init__(
        self
    ):
        self.document = Document()  

    def create_table(self, table_data, table_name):
        if not table_data or not isinstance(table_data[0], dict):
            return None
        
        table_width = len(table_data[0].keys())
        table_height = len(table_data)

        table = self.document.add_table(rows=table_height, cols=table_width)

        data = []

        column_headers = list(table_data[0].keys())
        data.append(column_headers)

        for row in table_data:
            if isinstance(row, dict):
                data.append(list(row.values()))

        for row_index, row_data in enumerate(data):
            for col_index, cell_data in enumerate(row_data):
                table.cell(row_index, col_index).text = cell_data

        logging.info(f"Created table for {table_name}")



    def format_table(self, analysis: Analysis):
        self.document.add_heading(f"{analysis.prompt_obj.get('display_name')}", level=2)
        self.document.add_paragraph(f"{analysis.prompt_obj.get('description')}")
        self.document.add_paragraph(f"{analysis.response.get('analysis', '')[:1900]}")

        table_data = analysis.response.get("table", [])
        self.create_table(table_data, analysis.prompt_obj.get('display_name'))

        logging.info(f"Formatted table for {analysis.prompt_obj.get('display_name')}")
 
    
    def format_analysis(self, analysis: Analysis):
        # Headings
        self.document.add_heading(f"{analysis.prompt_obj.get('display_name')}", level=2)

        # Description
        self.document.add_paragraph(f"{analysis.prompt_obj.get('description')}")

        # Analysis
        self.document.add_paragraph(f"{analysis.response.get('analysis')[0:1900]}")

        # Dot Pointgs
        for dot_point in analysis.response.get("dot_point_summary"):
            for key, item in dot_point.items():
                dot1 = self.document.add_paragraph('', style='ListBullet').add_run(key)
                dot1.bold = True
                dot1.underline = True

                self.document.add_paragraph(item, style='ListBullet2')

    def format_timeline(self, timeline: Analysis):
        # Headings
        self.document.add_heading(f"{timeline.prompt_obj.get('display_name')}", level=2)

        # Description
        self.document.add_paragraph(f"{timeline.prompt_obj.get('description')}")

        # Timeline Items
        for timeline_item in timeline.response.get("timeline"):
            for key, item in timeline_item.items():
                self.document.add_paragraph(key, style='ListBullet')
                self.document.add_paragraph(item, style='ListBullet2')

    def format_cost_value(self, cost_value: Analysis):
        # Headings
        self.document.add_heading(f"{cost_value.prompt_obj.get('display_name')}", level=2)

        # Description
        self.document.add_paragraph(f"{cost_value.prompt_obj.get('description')}")

        # cost_value Items
        cost_value_items = []
        for cost_value_item in cost_value.response.get("cost_value"):
            for key, item in cost_value_item.items():
                self.document.add_paragraph(key, style='ListBullet')
                self.document.add_paragraph(item, style='ListBullet2')

    def create_docx_from_analysis(self, proposal_name: str, analysis_list: list[Analysis]):
        logging.info(f"[DocxOperator] Creating file for the analysis result of {proposal_name}")

        australia_tz = pytz.timezone("Australia/Sydney")
        str_now = datetime.now(australia_tz).strftime("%d %b %Y - %I:%M:%S %p")
        doc_title = self.document.add_heading('', level=1)
        doc_title_run1 = doc_title.add_run('Analysis: ')
        doc_title_run2 = doc_title.add_run(f'{proposal_name}')
        doc_title_run2.italic = True

        doc_sub_title = self.document.add_paragraph('')
        doc_sub_title_run1 = doc_sub_title.add_run('Date Analysed: ')
        doc_sub_title_run2 = doc_sub_title.add_run(f'{str_now}')
        doc_sub_title_run2.bold = True
        doc_sub_title_run2.italic = True

        for analysis in analysis_list:
            if not isinstance(analysis, Analysis):
                logging.warning(f"Skipping invalid analysis object: {analysis}")
                continue

            if "table" in analysis.response:
                logging.info(f"[DocxOperator] Formatting table for the analysis result of {proposal_name}")

                self.format_table(analysis)
            elif "analysis" in analysis.response:
                logging.info(f"[DocxOperator] Formatting analysis for the analysis result of {proposal_name}")

                self.format_analysis(analysis)
            elif "timeline" in analysis.response:
                logging.info(f"[DocxOperator] Formatting timeline for the analysis result of {proposal_name}")

                self.format_timeline(analysis)
            elif "cost_value" in analysis.response:
                logging.info(f"[DocxOperator] Formatting cost value for the analysis result of {proposal_name}")

                self.format_cost_value(analysis)
            else:
                logging.warning(f"Unknown analysis type for {analysis.prompt_obj.get('display_name')}")

        docx_stream = io.BytesIO()
        self.document.save(docx_stream)
        docx_stream.seek(0)
        docx_bytes = docx_stream.getvalue()

        logging.info(f"[DocxOperator] File saved for the analysis result of {proposal_name}")


        analysed_proposal_filename = f"[Analysed] {proposal_name}.docx"

        return docx_bytes, analysed_proposal_filename, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'